import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Создание документа
doc = docx.Document()

# Настройка полей
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Цветовая палитра
COLOR_PRIMARY = RGBColor(0, 80, 179)    # Deep Blue
COLOR_SECONDARY = RGBColor(85, 85, 85)  # Gray
COLOR_TEXT = RGBColor(51, 51, 51)       # Dark Charcoal

# Стилизация текста по умолчанию
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = COLOR_TEXT

# Функция для добавления красивых заголовков
def add_heading(text, level, space_before=12, space_after=6):
    heading = doc.add_heading(level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
    else:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
    return heading

# --- ТИТУЛЬНЫЙ ЛИСТ ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(100)
title_run = title_p.add_run("ТЕХНИЧЕСКИЙ ДИЗАЙН-ПРОЕКТ\n")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run("Однокомнатная квартира-студия | 30.4 м²\nВысота потолков: 2.40 м\n\n")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = COLOR_SECONDARY

doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_p.paragraph_format.space_before = Pt(150)
desc_run = desc_p.add_run("Рабочая документация и техническое задание для строительной бригады\n2026 г.")
desc_run.font.size = Pt(10)
desc_run.font.italic = True

doc.add_page_break()

# --- ВВЕДЕНИЕ ---
add_heading("1. Общая информация и особенности объекта", level=1)
intro_p = doc.add_paragraph(
    "Данный проект разработан на основе фактических замеров, видеосъемки объекта и детальных рукописных схем заказчика. "
    "Главная цель проекта — обеспечить строительную бригаду точными размерами, планами расстановки сантехники и электрики "
    "для проведения качественного капитального ремонта."
)
intro_p.paragraph_format.space_after = Pt(12)

# Блок предупреждения
warn_table = doc.add_table(rows=1, cols=1)
warn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = warn_table.cell(0, 0)
cell.text = (
    "Внимание для прораба: Высота потолков в черновом состоянии составляет 2.40 м. "
    "Категорически не допускается создание многоуровневых гипсокартонных конструкций, "
    "существенно занижающих высоту. Рекомендуется использовать тонкий натяжной потолок (матовый) "
    "или подготовить плиту перекрытия под покраску. Проводку коммуникаций по полу вести в тонкой гофре, "
    "минимизируя толщину будущей стяжки/наливного пола."
)
for paragraph in cell.paragraphs:
    paragraph.runs[0].font.italic = True
    paragraph.runs[0].font.size = Pt(10)
    paragraph.runs[0].font.color.rgb = RGBColor(180, 0, 0)

# --- ДЕМОНТАЖ И ПЛАНИРОВКА ---
add_heading("2. Демонтаж и геометрия стен", level=1)
doc.add_paragraph("Перед началом монтажных работ строителям необходимо выполнить:")
bullets = [
    "Полный демонтаж старой плитки со стен и пола в санузле.",
    "Снятие напольного покрытия (ламинат, плитка) и удаление старой стяжки до плиты перекрытия.",
    "Демонтаж дверных коробок, старой сантехники, кухонной мебели и техники.",
    "Очистка стен от старых обоев, шпаклевки и фартука."
]
for bullet in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bullet)

# --- ПЛАНИРОВКА И МЕБЕЛЬ ---
add_heading("3. Спецификация мебели и габаритов по зонам", level=1)
doc.add_paragraph("Ниже приведены фиксированные габариты ключевых зон квартиры согласно чертежам:")

# Таблица мебели
table_data = [
    ("Зона", "Элемент", "Габариты (Ш х Г, см)", "Примечания"),
    ("Прихожая", "Встроенный шкаф", "178 х 60", "Монтаж от угла"),
    ("Прихожая", "Вешалка", "58 х 50", "Открытая зона у входа"),
    ("Санузел", "Душевой поддон", "90 х 90", "Со скошенным углом"),
    ("Санузел", "Унитаз (инсталляция)", "48 (длина)", "Короб стояка/инсталляции 30 см"),
    ("Санузел", "Тумба с раковиной", "69 х 30", "Слева от входа"),
    ("Санузел", "Стиральная машина", "50 х 50", "Слева от входа в углу"),
    ("Кухня", "Линейный гарнитур", "241 (общая длина)", "Разбивка модулей: 40-60-20-60-60"),
    ("Кухня", "Холодильник", "58 х 58", "Отдельностоящий в углу"),
    ("Спальня", "Кровать", "213 х 102", "По правой стене"),
    ("Гостиная", "Диван", "180 х 130", "В углу у выхода на террасу")
]

table = doc.add_table(rows=len(table_data), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row in enumerate(table_data):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = val
        # Стилизация заголовка таблицы
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.fill.solid()
            cell.fill.fore_color.rgb = COLOR_PRIMARY

# --- КУХНЯ ---
add_heading("4. Конфигурация кухонного гарнитура", level=2)
doc.add_paragraph("Общая длина кухонной линии по чертежу составляет 240 см (+1 см зазор):")
kitchen_steps = [
    "Секция 1: Бутылочница/карго — 40 см.",
    "Секция 2: Варочная панель и духовой шкаф — 60 см (требуется силовой вывод).",
    "Секция 3: Бутылочница/карго — 20 см.",
    "Секция 4: Модуль с выдвижными ящиками (рабочая зона) — 60 см.",
    "Секция 5: Шкаф под мойку — 60 см (выводы воды и канализации внутри секции)."
]
for step in kitchen_steps:
    doc.add_paragraph(style='List Bullet').add_run(step)

# --- САНТЕХНИКА ---
add_heading("5. Инженерный план сантехнических узлов", level=1)
doc.add_paragraph("При разводке водопровода и канализации строителям необходимо строго придерживаться следующих осей:")
plumb_steps = [
    "Душ (90х90 см): Выводы ГВС/ХВС на высоте 100 см от чистого пола по центру душевой зоны. Трап в полу со сверхнизким сифоном.",
    "Унитаз: Инсталляция крепится к капитальной стене. Выпуск канализации стандартный (высота 18-22 см по центру).",
    "Раковина (69 см): Выводы ХВС/ГВС на высоте 55 см от пола строго по центру раковины. Канализационный вывод диаметром 32 мм на высоте 50-53 см.",
    "Стиральная машина (50х50 см): Кран подачи холодной воды на высоте 60 см. Слив в канализацию через сифон с обратным клапаном (h=60 см)."
]
for step in plumb_steps:
    doc.add_paragraph(style='List Bullet').add_run(step)

# --- ЭЛЕКТРИКА ---
add_heading("6. План электрики и слаботочных сетей (Спецификация розеток)", level=1)
doc.add_paragraph("Все розетки монтируются строго в соответствии со следующей спецификацией высот и привязок:")

elect_data = [
    ("Зона", "Назначение розетки/вывода", "Тип", "Рекомендуемая высота (h, см)", "Кол-во"),
    ("Прихожая", "Входная зона / пылесос", "Обычная", "30", "1"),
    ("Прихожая", "Зона у шкафа", "Обычная", "30", "1"),
    ("Санузел", "Стиральная машина", "Силовой вывод / Влагозащищенная", "60", "1"),
    ("Санузел", "Около зеркала (фен, щетки)", "Влагозащищенная IP44", "100", "1"),
    ("Кухня", "Электроплита / Духовка", "Силовая розетка (32А)", "10-15 (в цоколе)", "1"),
    ("Кухня", "Рабочая зона (фартук)", "Блок обычных розеток", "105", "2"),
    ("Кухня", "Холодильник", "Силовой вывод / Обычная", "30", "1"),
    ("Спальня", "Левая прикроватная", "Обычная", "60", "1"),
    ("Спальня", "Правая прикроватная", "Обычная", "60", "1"),
    ("Гостиная", "Зона дивана (торшер/зарядки)", "Обычная", "30", "1"),
    ("Гостиная", "Рабочий стол", "Блок обычных розеток", "85 (над столом)", "2"),
    ("Комната", "ТВ-блок (за телевизором)", "Обычные + интернет RJ45", "110", "3"),
    ("Комната", "Кондиционер", "Силовой вывод (кабель)", "210", "1")
]

el_table = doc.add_table(rows=len(elect_data), cols=5)
el_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row in enumerate(elect_data):
    for j, val in enumerate(row):
        cell = el_table.cell(i, j)
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.fill.solid()
            cell.fill.fore_color.rgb = COLOR_PRIMARY

# Сохранение файла
doc.save("Design_Project_Studio_30_4m.docx")
print("Документ Word успешно создан!")